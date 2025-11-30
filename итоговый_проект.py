import sys
import subprocess as sp
from datetime import datetime
from docx import Document
from PyQt5.QtWidgets import *
from PyQt5.QtCore import QThread, pyqtSignal

class CheckThread(QThread):
    result_signal = pyqtSignal(str)
    finished_signal = pyqtSignal(list)

    def __init__(self, hosts):
        super().__init__()
        self.hosts = hosts

    def run(self):
        results = []
        for host in self.hosts:
            result = self.check_host(host)
            results.append(result)
            self.result_signal.emit(f"Проверен {host}: {result['status']}")
        self.finished_signal.emit(results)

    def check_host(self, host):
        result = sp.run(['ping', '-n', '2', host], stdout=sp.PIPE, stderr=sp.PIPE, creationflags=sp.CREATE_NO_WINDOW)
        status = 'Доступен' if result.returncode == 0 else 'Нет доступа'
        return {
            'host': host,
            'status': status,
            'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        }

class NetworkCheckApp(QWidget):
    def __init__(self):
        super().__init__()
        self.setup_ui()
        self.results_data = []

    def setup_ui(self):
        self.setWindowTitle('Проверка сети')
        self.resize(500, 400)
        layout = QVBoxLayout()
        
        layout.addWidget(QLabel('Введите адреса через пробел:'))
        self.host_input = QLineEdit()
        layout.addWidget(self.host_input)
        
        self.check_btn = QPushButton('Проверить')
        self.check_btn.clicked.connect(self.start_check)
        layout.addWidget(self.check_btn)
        
        self.report_btn = QPushButton('Создать отчет')
        self.report_btn.clicked.connect(self.create_report)
        self.report_btn.setEnabled(False)
        layout.addWidget(self.report_btn)
        
        layout.addWidget(QLabel('Результаты:'))
        self.results_list = QListWidget()
        layout.addWidget(self.results_list)
        
        layout.addWidget(QLabel('Лог:'))
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        layout.addWidget(self.log_text)
        
        self.setLayout(layout)

    def valid(self, host):
        host = host.strip()
        if not host:
            return False
            
        if host.endswith('.com') or host.endswith('.ru') or host.endswith('.org') or host.endswith('.net'):
            return True
        
        if host.count('.') == 3:
            parts = host.split('.')
            if len(parts) == 4:
                for part in parts:
                    try:
                        num = int(part)
                        if num < 0 or num > 255:
                            return False
                    except ValueError:
                        return False
                return True
        return False

    def start_check(self):
        text = self.host_input.text().strip()
        if not text:
            QMessageBox.warning(self, 'Ошибка', 'Введите адреса')
            return
        hosts = [h.strip() for h in text.split() if self.valid(h.strip())]
        if not hosts:
            QMessageBox.warning(self, 'Ошибка', 'Нет правильных адресов')
            return
        
        self.results_list.clear()
        self.log_text.clear()
        self.results_data = []
        
        self.thread = CheckThread(hosts)
        self.thread.result_signal.connect(self.log_text.append)
        self.thread.finished_signal.connect(self.done_check)
        
        self.check_btn.setEnabled(False)
        self.log_text.append(f'Проверяем {len(hosts)} хостов...')
        self.thread.start()

    def done_check(self, results):
        self.results_data = results
        self.check_btn.setEnabled(True)
        self.report_btn.setEnabled(True)
        for item in results:
            self.results_list.addItem(f"{item['host']} - {item['status']}")
        available = sum(1 for item in results if item['status'] == 'Доступен')
        self.log_text.append(f'Доступно: {available} из {len(results)}')

    def create_report(self):
        if not self.results_data:
            return
        
        doc = Document()
        doc.add_heading('Отчёт проверки сети', 0)
        doc.add_paragraph(f"Дата: {datetime.now().strftime('%d-%m-%Y %H:%M:%S')}")
        doc.add_paragraph(f'Всего узлов: {len(self.results_data)}')
        
        table = doc.add_table(rows=1, cols=3)
        h = table.rows[0].cells
        h[0].text, h[1].text, h[2].text = 'Адрес', 'Статус', 'Время'
        
        for item in self.results_data:
            r = table.add_row().cells
            r[0].text, r[1].text, r[2].text = item['host'], item['status'], item['time']
        
        available = sum(1 for item in self.results_data if item['status'] == 'Доступен')
        doc.add_paragraph(f'Доступно: {available} из {len(self.results_data)}')
        
        filename = f"отчет_{datetime.now().strftime('%d-%m-%Y_%H-%M')}.docx"
        doc.save(filename)
        
        self.log_text.append(f'Отчёт: {filename}')
        QMessageBox.information(self, 'Готово', f'Отчет: {filename}')

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = NetworkCheckApp()
    window.show()
    sys.exit(app.exec_())